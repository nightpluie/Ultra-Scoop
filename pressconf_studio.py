#!/usr/bin/env python3
from __future__ import annotations
"""
ULTRA SCOOP  v3.4.1
Mixer Layout: 輸入 → 處理 → 成稿
v3.4: 介面簡化——設定集中至「⚙ 設定」對話框、輸入來源改為分頁切換、
      查核結果點擊跳轉、錄音備份自動清理
"""

import sys, os, re, datetime, threading, json, math, time
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter.scrolledtext import ScrolledText

try:
    import customtkinter as ctk
except ImportError:
    raise SystemExit(
        "缺少 customtkinter，請執行：pip install customtkinter")

# ── paths ─────────────────────────────────────────────────────────────────────
SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
WHISPER_DIR = SCRIPT_DIR
SKILL_PATH  = os.path.join(SCRIPT_DIR, "skills", "report-tcy", "report-tcy.md")
CONFIG_PATH = os.path.join(SCRIPT_DIR, "config.json")

# ── ffmpeg：系統沒裝就改用 pip 自帶版本，確保轉錄功能可用 ──────────────────────
try:
    sys.path.insert(0, SCRIPT_DIR)
    from ffmpeg_setup import ensure_ffmpeg_on_path
    ensure_ffmpeg_on_path()
except Exception:
    pass

# ── optional deps ─────────────────────────────────────────────────────────────
def _has(m):
    try: __import__(m); return True
    except ImportError: return False

HAS_SOUNDDEVICE = _has("sounddevice")
HAS_MLX_WHISPER = _has("mlx_whisper")

# ── 拆分模組（v3.4）───────────────────────────────────────────────────────────
from app_theme import *
from text_utils import (parse_single_output, parse_dual_output, parse_file,
                        strip_markdown, extract_numbers,
                        HAS_PDFPLUMBER, HAS_DOCX)
from claude_api import (HAS_ANTHROPIC, claude_generate, claude_check,
                        claude_correct, claude_translate)
from ui_helpers import _add_copy_menu, _add_entry_menu, _dark_text, _select_all

try:
    sys.path.insert(0, WHISPER_DIR)
    sys.path.insert(0, SCRIPT_DIR)
    from gui_whisper_live import (
        LiveTranscriber, WHISPER_MODELS, LANGUAGES, SPEAKER_PALETTE
    )
    HAS_LIVE = True
except Exception:
    HAS_LIVE        = False
    WHISPER_MODELS  = ["tiny", "base", "small", "medium", "large-v2", "large-v3", "large-v3-turbo", "turbo"]
    LANGUAGES       = ["auto", "zh", "en"]
    SPEAKER_PALETTE = ["#5B8DB8", "#8E78C0"]

try:
    from phone_mic_server import (PhoneMicTranscriber, _get_local_ip,
                                  make_qr_image_tk, ensure_server_running)
    HAS_PHONE_MIC = True
except Exception:
    HAS_PHONE_MIC = False

# ── config ────────────────────────────────────────────────────────────────────
def load_config() -> dict:
    try:
        with open(CONFIG_PATH, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}

def save_config(data: dict):
    try:
        cfg = load_config()
        cfg.update(data)
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

# ══════════════════════════════════════════════════════════════════════════════
class PressConfStudio:
    def __init__(self, root: ctk.CTk):
        self.root = root
        root.title("ULTRA SCOOP  v3.4.1")
        sw, sh = root.winfo_screenwidth(), root.winfo_screenheight()
        w = min(int(sw * 0.95), 1900)
        h = min(sh - 110, 1020)
        x = (sw - w) // 2
        y = 30
        root.geometry(f"{w}x{h}+{x}+{y}")
        root.minsize(900, 600)

        self._transcriber  = None
        self._log_file     = None
        self._device_map: dict[str, int] = {}
        self._files: list[tuple[str, str]] = []
        self._generating   = False
        self._checking     = False
        self._audio_file   = ""
        self._audio_busy   = False
        self._last_scene   = ""   # 最近一次錄音的記者會名稱，供校正階段使用

        # wave animation state
        self._wave_anim_id  = None
        self._wave_frame    = 0
        self._wave_active   = False

        # record button pulse animation
        self._rec_anim_id   = None
        self._rec_pulse_on  = False

        # sidebar mode toggle
        self._sidebar_mode = tk.BooleanVar(value=False)

        # phone mic mode
        self._phone_mode = tk.BooleanVar(value=False)
        self._qr_photo   = None

        # remember preferences
        self._remember_author = tk.BooleanVar(value=False)
        self._remember_key    = tk.BooleanVar(value=False)

        # settings (managed via「⚙ 設定」dialog)
        self._wmodel_var = tk.StringVar(value="large-v3-turbo")
        self._lang_var   = tk.StringVar(value="zh")
        self._gen_model  = tk.StringVar(value="claude-sonnet-4-6")
        self._apikey_var = tk.StringVar()
        self._settings_win = None
        self._skill_lbl     = None
        self._apikey_entry  = None

        # skill path
        self._skill_path_var = tk.StringVar(value=SKILL_PATH)

        # last fact-check issues
        self._last_check_issues: list = []

        # api key visibility
        self._apikey_visible = False

        # col3 tab state ("main" / "side")
        self._col3_active = "main"

        self._build_ui()
        self._refresh_devices()
        self._bind_shortcuts()

        cfg = load_config()
        if cfg.get("author"):
            self._author_var.set(cfg["author"])
            self._remember_author.set(True)
        if cfg.get("api_key"):
            self._apikey_var.set(cfg["api_key"])
            self._remember_key.set(True)
        if cfg.get("skill_path"):
            self._skill_path_var.set(cfg["skill_path"])
        self._skill_path_var.trace_add("write", lambda *_: self._update_skill_label())
        self._update_skill_label()   # 確保 config 載入後 label 正確顯示

        self._apikey_var.trace_add("write", lambda *_: self._update_key_label())
        self._update_key_label()

        # 首次使用：尚未設定 API Key 時自動開啟設定視窗引導
        if not (self._apikey_var.get().strip()
                or os.environ.get("ANTHROPIC_API_KEY")):
            self.root.after(600, self._first_run_hint)

        # 錄音備份清理檢查（避免 rec/ 無限增長佔滿磁碟）
        self.root.after(1500, self._check_rec_cleanup)

    def _first_run_hint(self):
        self._status("尚未設定 API Key — 請在「⚙ 設定」填入後即可使用生成功能")
        self._open_settings()

    # ── rec/ 錄音備份清理 ────────────────────────────────────────────────────
    REC_KEEP_DAYS = 30

    def _check_rec_cleanup(self):
        """啟動時檢查 rec/，詢問是否刪除超過保留天數的錄音備份。"""
        rec_dir = os.path.join(SCRIPT_DIR, "rec")
        if not os.path.isdir(rec_dir):
            return
        cutoff = time.time() - self.REC_KEEP_DAYS * 86400
        old_files, total_bytes = [], 0
        for dirpath, _, names in os.walk(rec_dir):
            for fn in names:
                if not fn.lower().endswith(".wav"):
                    continue
                path = os.path.join(dirpath, fn)
                try:
                    st = os.stat(path)
                except OSError:
                    continue
                if st.st_mtime < cutoff:
                    old_files.append(path)
                    total_bytes += st.st_size
        if not old_files:
            return
        mb = total_bytes / 1048576
        if not messagebox.askyesno(
                "清理錄音備份",
                f"rec/ 資料夾有 {len(old_files)} 個超過 {self.REC_KEEP_DAYS} 天"
                f"的錄音備份（約 {mb:.0f} MB）。\n\n"
                "要刪除這些舊檔嗎？\n"
                "（逐字稿文字備份存放在 log/，不受影響）"):
            return
        removed = 0
        for path in old_files:
            try:
                os.remove(path)
                removed += 1
            except OSError:
                pass
        self._status(f"已清理 {removed} 個舊錄音備份（釋出約 {mb:.0f} MB）")

    # ── keyboard shortcuts ────────────────────────────────────────────────────
    def _bind_shortcuts(self):
        self.root.bind_all("<Command-g>",      lambda e: self._generate())
        self.root.bind_all("<Command-k>",      lambda e: self._verify())
        self.root.bind_all("<Command-s>",      lambda e: self._save_txt())
        self.root.bind_all("<Command-Return>", lambda e: self._generate())
        self.root.bind_all("<Command-comma>",  lambda e: self._open_settings())
        # macOS：中文輸入法啟用或 Caps Lock 開啟時，keysym 會變形
        # （'C'、'??' 等），上述依 keysym 比對的綁定全部失效。
        # 此萬用攔截層在 keysym 變形時改用硬體鍵碼／char 判斷按鍵，
        # 再走與選單列相同的 <<Copy>>/<<Paste>> 虛擬事件路徑。
        self.root.bind_all("<Command-KeyPress>", self._cmd_key_fallback)

    # 實體按鍵碼（不受輸入法／大小寫影響）；Tk 9 aqua：keycode >> 24
    _MAC_VK = {0: "a", 1: "s", 5: "g", 7: "x", 8: "c", 9: "v",
               40: "k", 36: "return", 43: "comma"}
    # keysym 正常時交由原綁定處理，避免重複觸發
    _HANDLED_KEYSYMS = {"a", "c", "v", "x", "g", "k", "s", "comma", "Return"}

    def _vk_key(self, keycode: int):
        if not keycode:
            return None
        if (keycode >> 16) & 0x80:
            # Tk 9 aqua：keycode = (實體鍵碼 << 24) | 0x80____ | 字元碼。
            # 實體鍵碼 0（A 鍵）需以低位元組為 ASCII 字母佐證，
            # 排除修飾鍵本身（如 0x80F8FE，實體鍵碼也是 0）的誤判。
            vk, low = keycode >> 24, keycode & 0xFFFF
            if vk in self._MAC_VK and (vk or 0x41 <= low <= 0x7A):
                return self._MAC_VK[vk]
            return None
        vk16 = keycode >> 16          # Tk 8.6 aqua：keycode = 實體鍵碼 << 16
        if vk16 and vk16 in self._MAC_VK:
            return self._MAC_VK[vk16]
        return self._MAC_VK.get(keycode)

    def _cmd_key_fallback(self, event):
        if event.keysym in self._HANDLED_KEYSYMS:
            return  # 正常路徑，原綁定已處理
        ks, ch = event.keysym, event.char
        if len(ks) == 1 and ks.isascii() and ks.isalpha():
            key = ks.lower()                      # Caps Lock／Shift："C"
        elif ch and len(ch) == 1 and ch.isascii() and ch.isalpha():
            key = ch.lower()                      # keysym 壞掉但 char 正常
        else:
            key = self._vk_key(event.keycode)     # 輸入法：用硬體鍵碼
        if key not in self._HANDLED_KEYSYMS and key != "return":
            return
        if key in ("c", "v", "x", "a"):
            w = self.root.focus_get()
            if w is None:
                return "break"
            if key == "a":
                if isinstance(w, tk.Text):
                    _select_all(w)
                elif hasattr(w, "select_range"):
                    w.select_range(0, tk.END)
            else:
                event_name = {"c": "<<Copy>>", "v": "<<Paste>>",
                              "x": "<<Cut>>"}[key]
                w.event_generate(event_name)
        elif key in ("g", "return"):
            self._generate()
        elif key == "k":
            self._verify()
        elif key == "s":
            self._save_txt()
        elif key == "comma":
            self._open_settings()
        return "break"

    # ── UI building ───────────────────────────────────────────────────────────
    def _build_ui(self):
        self._build_topbar()
        self._build_statusbar()
        self._build_main()

    def _build_topbar(self):
        bar = ctk.CTkFrame(self.root, fg_color=BG_SURFACE, height=56,
                           corner_radius=0)
        bar.pack(fill="x")
        bar.pack_propagate(False)

        ctk.CTkLabel(bar, text="ULTRA",
                     font=("Arial Black", 26),
                     text_color="#FFFFFF").pack(side="left", padx=(16, 3))
        ctk.CTkLabel(bar, text="SCOOP",
                     font=("Arial Black", 26),
                     text_color=ACCENT_AMBER).pack(side="left", padx=(0, 6))
        ctk.CTkLabel(bar, text="v3.4.1",
                     font=("Menlo", 9), text_color=TEXT_DIM
                     ).pack(side="left", padx=(0, 16), pady=(10, 0))

        # separator
        sep = ctk.CTkFrame(bar, fg_color=BORDER, width=1, height=24,
                           corner_radius=0)
        sep.pack(side="left", padx=8)

        # pipeline indicator
        for label, color in [
            ("1 輸入", COL_INPUT),
            ("-->",    TEXT_DIM),
            ("2 處理", COL_PROCESS),
            ("-->",    TEXT_DIM),
            ("3 成稿", COL_OUTPUT),
        ]:
            ctk.CTkLabel(bar, text=label,
                         font=("PingFang TC", 10),
                         text_color=color).pack(side="left", padx=3)

        # shortcuts hint
        ctk.CTkLabel(bar, text="Cmd+G 生成  Cmd+K 查核  Cmd+S 存檔",
                     font=FT_SM, text_color=TEXT_DIM
                     ).pack(side="left", padx=(16, 0))

        # settings button + API key status (right side)
        ctk.CTkButton(bar, text="⚙ 設定", width=72, height=30,
                      font=FT_SM, fg_color=BG_INPUT, hover_color=BORDER,
                      text_color=TEXT_PRI, border_width=1, border_color=BORDER,
                      command=self._open_settings
                      ).pack(side="right", padx=(0, 16))
        self._topbar_key_lbl = ctk.CTkLabel(bar, text="", font=FT_SM)
        self._topbar_key_lbl.pack(side="right", padx=12)

    def _build_statusbar(self):
        bar = ctk.CTkFrame(self.root, fg_color=BG_SURFACE, height=26,
                           corner_radius=0)
        bar.pack(fill="x", side="bottom")
        bar.pack_propagate(False)

        self._status_var = tk.StringVar(value="就緒")
        ctk.CTkLabel(bar, textvariable=self._status_var,
                     font=FT_SM, text_color=TEXT_PRI,
                     anchor="w").pack(side="left", padx=10)

        ctk.CTkLabel(bar, text="Designed by Chen-Yu TANG",
                     font=FT_SM, text_color=TEXT_DIM
                     ).pack(side="right", padx=10)

        missing = [p for p, ok in [
            ("anthropic",      HAS_ANTHROPIC),
            ("sounddevice",    HAS_SOUNDDEVICE),
            ("mlx-whisper", HAS_MLX_WHISPER),
            ("pdfplumber",     HAS_PDFPLUMBER),
        ] if not ok]
        if missing:
            ctk.CTkLabel(bar, text=f"未安裝：{', '.join(missing)}",
                         font=FT_SM, text_color=WARN_FG
                         ).pack(side="right", padx=10)

    # ── 3-column main layout ─────────────────────────────────────────────────
    def _build_main(self):
        paned = tk.PanedWindow(self.root, orient=tk.HORIZONTAL,
                               bg=BG, sashwidth=5, sashrelief="flat",
                               bd=0, sashpad=0)
        paned.pack(fill="both", expand=True, padx=4, pady=4)

        col1 = ctk.CTkFrame(paned, fg_color=BG_SURFACE, corner_radius=0,
                            border_width=1, border_color=BORDER)
        paned.add(col1, width=460, minsize=300, stretch="always")

        col2 = ctk.CTkFrame(paned, fg_color=BG_SURFACE, corner_radius=0,
                            border_width=1, border_color=BORDER)
        paned.add(col2, width=520, minsize=340, stretch="always")

        col3 = ctk.CTkFrame(paned, fg_color=BG_SURFACE, corner_radius=0,
                            border_width=1, border_color=BORDER)
        paned.add(col3, width=480, minsize=300, stretch="always")

        self._build_col_input(col1)
        self._build_col_process(col2)
        self._build_col_output(col3)

    def _channel_header(self, parent, label, accent_color):
        """DAW-style channel strip header with colored accent bar."""
        accent = ctk.CTkFrame(parent, fg_color=accent_color, height=3,
                              corner_radius=0)
        accent.pack(fill="x")
        accent.pack_propagate(False)

        hdr = ctk.CTkFrame(parent, fg_color=BG_PANEL, height=32, corner_radius=0)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        ctk.CTkLabel(hdr, text=label, font=FT_BIG, text_color=TEXT_PRI,
                     anchor="w").pack(side="left", padx=12, pady=4)
        return hdr

    def _make_combo(self, parent, variable, values, width=120):
        """統一樣式的唯讀下拉選單。"""
        return ctk.CTkComboBox(
            parent, variable=variable, state="readonly", width=width,
            values=values, font=FT_SM, dropdown_font=FT_SM,
            fg_color=BG_INPUT, border_color=BORDER,
            button_color=BG_PANEL, button_hover_color=BORDER_LT,
            dropdown_fg_color=BG_PANEL, dropdown_hover_color=BORDER_LT,
            text_color=TEXT_PRI)

    def _ghost_btn(self, parent, text, command,
                   width=40, height=26, text_color=TEXT_SEC):
        """統一樣式的次要（描邊）按鈕。"""
        return ctk.CTkButton(
            parent, text=text, width=width, height=height,
            font=FT_SM, fg_color=BG_INPUT, hover_color=BORDER,
            text_color=text_color, border_width=1, border_color=BORDER,
            command=command)

    # ══════════════════════════════════════════════════════════════════════════
    #  COLUMN 1 — 輸入素材
    # ══════════════════════════════════════════════════════════════════════════
    def _build_col_input(self, parent):
        self._channel_header(parent, "輸入素材", COL_INPUT)
        body = ctk.CTkFrame(parent, fg_color="transparent")
        body.pack(fill="both", expand=True, padx=6, pady=6)

        # ── Source selector：一次只顯示一種輸入方式 ────────────────────────
        self._active_src = "現場錄音"
        self._src_var = tk.StringVar(value=self._active_src)
        self._src_seg = ctk.CTkSegmentedButton(
            body, values=["現場錄音", "手機麥克風", "音檔轉錄", "附件"],
            variable=self._src_var, command=self._switch_source,
            font=FT_SM, height=30,
            fg_color=BG_PANEL,
            selected_color=COL_INPUT, selected_hover_color=COL_INPUT_HOVER,
            unselected_color=BG_PANEL, unselected_hover_color=BORDER,
            text_color=TEXT_PRI)
        self._src_seg.pack(fill="x", pady=(0, 4))

        self._src_panel = ctk.CTkFrame(body, fg_color=BG_PANEL,
                                       corner_radius=6)
        self._src_panel.pack(fill="x", pady=(0, 4))

        self._src_frames = {
            "現場錄音":   self._build_src_live(self._src_panel),
            "手機麥克風": self._build_src_phone(self._src_panel),
            "音檔轉錄":   self._build_src_audio(self._src_panel),
            "附件":       self._build_src_files(self._src_panel),
        }
        self._src_frames[self._active_src].pack(fill="x")

        # ── Shared recording feedback ─────────────────────────────────────
        self._live_lbl = ctk.CTkLabel(body, text="", font=FT_SM,
                                      text_color=TEXT_SEC)
        self._live_lbl.pack(pady=(0, 2))
        self._wave_canvas = tk.Canvas(body, height=24, bg=BG_INPUT,
                                      highlightthickness=0)
        self._wave_canvas.pack(fill="x", pady=(0, 4))

        # ── Transcript toolbar ────────────────────────────────────────────
        tools = ctk.CTkFrame(body, fg_color="transparent")
        tools.pack(fill="x", pady=(0, 2))
        ctk.CTkLabel(tools, text="逐字稿", font=FT_SM,
                     text_color=TEXT_SEC).pack(side="left", padx=(2, 6))
        for label, cmd in [("貼上", self._paste_transcript),
                           ("清除", self._live_clear),
                           ("載入", self._load_transcript),
                           ("儲存", self._save_transcript)]:
            self._ghost_btn(tools, label, cmd).pack(side="left", padx=2)
        # AI 輔助小工具統一以紫色文字識別（規則 3）
        self._correct_btn = self._ghost_btn(
            tools, "修正錯字", self._correct_transcript,
            width=68, text_color=ACCENT_PURPLE)
        self._correct_btn.pack(side="left", padx=2)
        self._trans_btn = self._ghost_btn(
            tools, "翻成中文", self._translate_transcript,
            width=68, text_color=ACCENT_PURPLE)
        self._trans_btn.pack(side="left", padx=2)

        # ── Transcript text area ──────────────────────────────────────────
        self._transcript = _dark_text(body, height=6)
        self._transcript.pack(fill="both", expand=True)
        for i, c in enumerate(SPEAKER_PALETTE):
            self._transcript.tag_configure(f"spk{i}", foreground=c, font=FT_BOLD)
        self._transcript.tag_configure("body", foreground=TEXT_PRI, font=FT)
        _add_copy_menu(self._transcript)

    # ── 來源分頁：各面板 ──────────────────────────────────────────────────
    def _build_src_live(self, parent):
        f = ctk.CTkFrame(parent, fg_color="transparent")

        r_dev = ctk.CTkFrame(f, fg_color="transparent")
        r_dev.pack(fill="x", padx=8, pady=(8, 2))
        ctk.CTkLabel(r_dev, text="輸入裝置", font=FT_SM,
                     text_color=TEXT_SEC).pack(side="left")
        self._dev_var = tk.StringVar()
        self._dev_combo = self._make_combo(r_dev, self._dev_var, [], width=180)
        self._dev_combo.pack(side="left", fill="x", expand=True, padx=(6, 4))
        self._ghost_btn(r_dev, "重整", self._refresh_devices,
                        height=24).pack(side="left")

        btns_row = ctk.CTkFrame(f, fg_color="transparent")
        btns_row.pack(fill="x", padx=8, pady=(4, 8))
        self._start_btn = tk.Canvas(
            btns_row, width=42, height=42,
            bg=BG_PANEL, highlightthickness=1,
            highlightbackground=BORDER, cursor="hand2")
        self._start_btn.pack(side="left", padx=(0, 4))
        self._rec_circle = self._start_btn.create_oval(
            11, 11, 31, 31, fill=ACCENT_RED, outline="")
        self._start_btn.bind("<Button-1>", lambda e: self._live_start())

        self._stop_btn = tk.Canvas(
            btns_row, width=42, height=42,
            bg=BG_PANEL, highlightthickness=1,
            highlightbackground=BORDER)
        self._stop_btn.pack(side="left", padx=(0, 8))
        self._stop_square = self._stop_btn.create_rectangle(
            12, 12, 30, 30, fill=TEXT_DIM, outline="")
        self._stop_btn.bind("<Button-1>", lambda e: self._live_stop()
                            if self._wave_active else None)
        ctk.CTkLabel(btns_row, text="紅鈕開始錄音　方鈕停止",
                     font=FT_SM, text_color=TEXT_DIM).pack(side="left", padx=6)
        return f

    def _build_src_phone(self, parent):
        f = ctk.CTkFrame(parent, fg_color="transparent")

        top = ctk.CTkFrame(f, fg_color="transparent")
        top.pack(fill="x", padx=8, pady=(8, 0))
        ctk.CTkLabel(top, text="手機瀏覽器開啟：", font=FT_SM,
                     text_color=TEXT_SEC).pack(side="left")
        self._phone_url_var = tk.StringVar(value="")
        url_lbl = ctk.CTkLabel(top, textvariable=self._phone_url_var,
                               font=("Menlo", 9), text_color=ACCENT_BLUE,
                               cursor="hand2")
        url_lbl.pack(side="left")
        url_lbl.bind("<Button-1>", lambda e: self._copy_phone_url())
        ctk.CTkLabel(top, text="（點按複製）", font=FT_SM,
                     text_color=TEXT_DIM).pack(side="left", padx=4)

        self._qr_label = tk.Label(f, bg=BG_PANEL)
        self._qr_label.pack(pady=(4, 2))

        bottom = ctk.CTkFrame(f, fg_color="transparent")
        bottom.pack(fill="x", padx=8, pady=(0, 8))
        ctk.CTkLabel(bottom, text="掃描後在手機網頁按「開始」即可錄音",
                     font=FT_SM, text_color=TEXT_DIM).pack(side="left")
        self._ghost_btn(bottom, "停止錄音",
                        lambda: self._live_stop() if self._wave_active else None,
                        width=72, height=24).pack(side="right")
        return f

    def _build_src_audio(self, parent):
        f = ctk.CTkFrame(parent, fg_color="transparent")

        ar1 = ctk.CTkFrame(f, fg_color="transparent")
        ar1.pack(fill="x", padx=8, pady=(8, 2))
        self._audio_name_var = tk.StringVar(value="尚未選擇音檔")
        ctk.CTkLabel(ar1, textvariable=self._audio_name_var,
                     font=FT_SM, text_color=TEXT_SEC, anchor="w"
                     ).pack(side="left", fill="x", expand=True)
        self._ghost_btn(ar1, "選擇音檔", self._audio_browse,
                        width=70, height=24,
                        text_color=TEXT_PRI).pack(side="right")

        ar2 = ctk.CTkFrame(f, fg_color="transparent")
        ar2.pack(fill="x", padx=8, pady=(2, 8))
        self._audio_btn = ctk.CTkButton(
            ar2, text="開始轉錄", width=80, height=28,
            font=FT_SM, fg_color=COL_INPUT, hover_color=COL_INPUT_HOVER,
            text_color="#FFFFFF", command=self._audio_transcribe)
        self._audio_btn.pack(side="left")
        self._audio_prog = ctk.CTkProgressBar(
            ar2, mode="indeterminate", width=80, height=4,
            fg_color=BG_INPUT, progress_color=COL_INPUT)
        # 預設隱藏
        self._audio_lbl = ctk.CTkLabel(ar2, text="", font=FT_SM,
                                       text_color=TEXT_SEC)
        self._audio_lbl.pack(side="left", padx=4)
        ctk.CTkLabel(ar2, text="支援 MP3 / WAV / M4A / 影片檔",
                     font=FT_SM, text_color=TEXT_DIM).pack(side="right")
        return f

    def _build_src_files(self, parent):
        f = ctk.CTkFrame(parent, fg_color="transparent")

        drop = ctk.CTkButton(
            f, text="點擊上傳附件　PDF / Word / Excel / TXT", height=32,
            font=FT_SM, fg_color=BG_INPUT, hover_color=BORDER,
            text_color=TEXT_PRI, border_width=1, border_color=BORDER,
            command=self._file_browse)
        drop.pack(fill="x", padx=8, pady=(8, 2))

        self._flist = tk.Listbox(
            f, height=2, font=FT_SM,
            bg=BG_INPUT, fg=TEXT_PRI,
            relief="flat", bd=0,
            selectmode=tk.SINGLE, activestyle="none",
            selectbackground="#2A4A6A", selectforeground=TEXT_PRI,
            highlightthickness=0)
        self._flist.pack(fill="x", padx=8, pady=2)
        self._flist.bind("<Double-Button-1>", self._file_preview)

        fb = ctk.CTkFrame(f, fg_color="transparent")
        fb.pack(fill="x", padx=8, pady=(0, 8))
        for label, cmd in [("預覽", self._file_preview),
                           ("移除", self._file_remove),
                           ("清空", self._file_clear)]:
            self._ghost_btn(fb, label, cmd, height=22).pack(side="left", padx=2)
        ctk.CTkLabel(fb, text="附件會連同逐字稿一起送給 AI 寫稿",
                     font=FT_SM, text_color=TEXT_DIM).pack(side="right")
        return f

    def _switch_source(self, choice: str):
        prev = self._active_src
        if choice == prev:
            return
        if choice == "手機麥克風":
            if not HAS_PHONE_MIC:
                messagebox.showwarning("缺少套件",
                    "請先執行安裝程式安裝 flask 套件後再使用手機麥克風功能")
                self._src_var.set(prev)
                return
            if self._wave_active and not self._phone_mode.get():
                messagebox.showwarning("錄音中",
                    "請先停止現場錄音，再切換至手機麥克風")
                self._src_var.set(prev)
                return
        if prev == "手機麥克風":
            self._set_phone_mode(False)
        self._src_frames[prev].pack_forget()
        self._src_frames[choice].pack(fill="x")
        self._active_src = choice
        if choice == "手機麥克風":
            self._set_phone_mode(True)

    # ══════════════════════════════════════════════════════════════════════════
    #  COLUMN 2 — 處理（生成報導 / 事實查核 Tab）
    # ══════════════════════════════════════════════════════════════════════════
    def _build_col_process(self, parent):
        self._channel_header(parent, "處理", COL_PROCESS)
        body = ctk.CTkFrame(parent, fg_color="transparent")
        body.pack(fill="both", expand=True, padx=4, pady=4)
        self._build_gen_section(body)
        # ── 分隔線 ────────────────────────────────────────────────────────
        ctk.CTkFrame(body, fg_color=BORDER, height=1).pack(fill="x", pady=(8, 0))
        self._build_chk_section(body)

    def _build_gen_section(self, body):
        # ── Writing settings ──────────────────────────────────────────────
        settings = ctk.CTkFrame(body, fg_color=BG_PANEL, corner_radius=6)
        settings.pack(fill="x", pady=(0, 4))

        # Interviewee
        r1 = ctk.CTkFrame(settings, fg_color="transparent")
        r1.pack(fill="x", padx=8, pady=(6, 2))
        ctk.CTkLabel(r1, text="受訪者", font=FT_SM, text_color=TEXT_SEC,
                     width=50, anchor="w").pack(side="left")
        self._interviewee_var = tk.StringVar()
        e1 = ctk.CTkEntry(r1, textvariable=self._interviewee_var, font=FT_SM,
                          fg_color=BG_INPUT, border_color=BORDER,
                          text_color=TEXT_PRI)
        e1.pack(side="left", fill="x", expand=True, padx=(4, 0))
        _add_entry_menu(e1)

        # Main angle
        r2 = ctk.CTkFrame(settings, fg_color="transparent")
        r2.pack(fill="x", padx=8, pady=2)
        ctk.CTkLabel(r2, text="主稿角度", font=FT_SM, text_color=TEXT_SEC,
                     width=50, anchor="w").pack(side="left")
        self._main_angle_var = tk.StringVar()
        e2 = ctk.CTkEntry(r2, textvariable=self._main_angle_var, font=FT_SM,
                          fg_color=BG_INPUT, border_color=BORDER,
                          text_color=TEXT_PRI)
        e2.pack(side="left", fill="x", expand=True, padx=(4, 0))
        _add_entry_menu(e2)

        # Sidebar toggle
        self._sidebar_chk_row = r3 = ctk.CTkFrame(settings, fg_color="transparent")
        r3.pack(fill="x", padx=8, pady=(2, 8))
        ctk.CTkCheckBox(r3, text="同時產出配稿", font=FT_SM,
                        text_color=TEXT_SEC, variable=self._sidebar_mode,
                        fg_color=BORDER, hover_color=BORDER_LT,
                        checkmark_color=TEXT_PRI,
                        command=self._toggle_sidebar).pack(side="left")

        # Sidebar angle (hidden by default)
        self._sidebar_ctx_frame = ctk.CTkFrame(settings, fg_color="transparent")
        r4 = ctk.CTkFrame(self._sidebar_ctx_frame, fg_color="transparent")
        r4.pack(fill="x", padx=8, pady=(0, 8))
        ctk.CTkLabel(r4, text="配稿角度", font=FT_SM, text_color=TEXT_SEC,
                     width=50, anchor="w").pack(side="left")
        self._side_angle_var = tk.StringVar()
        e3 = ctk.CTkEntry(r4, textvariable=self._side_angle_var, font=FT_SM,
                          fg_color=BG_INPUT, border_color=BORDER,
                          text_color=TEXT_PRI)
        e3.pack(side="left", fill="x", expand=True, padx=(4, 0))
        _add_entry_menu(e3)

        # ── Generate button ───────────────────────────────────────────────
        self._gen_btn = ctk.CTkButton(
            body, text="生成報導", height=38, font=FT_BOLD,
            fg_color=COL_PROCESS, hover_color=COL_PROCESS_HOVER,
            text_color="#FFFFFF", command=self._generate)
        self._gen_btn.pack(fill="x", pady=(0, 2))

        self._gen_prog = ctk.CTkProgressBar(
            body, mode="indeterminate", height=4,
            fg_color=BG_INPUT, progress_color=COL_PROCESS)
        # 預設隱藏，生成時才顯示

        self._gen_lbl = ctk.CTkLabel(body, text="等待生成", font=FT_SM,
                                     text_color=TEXT_SEC)
        self._gen_lbl.pack()

    def _build_chk_section(self, body):
        # ── Stats panel ───────────────────────────────────────────────────
        stats_panel = ctk.CTkFrame(body, fg_color=BG_PANEL, corner_radius=6)
        stats_panel.pack(fill="x", pady=(4, 4))
        ctk.CTkLabel(stats_panel, text="數字出處統計", font=FT_SM,
                     text_color=TEXT_SEC).pack(padx=8, pady=(6, 2), anchor="w")
        self._stats_lbl = ctk.CTkLabel(
            stats_panel, text="尚未執行查核", font=FT_SM,
            text_color=TEXT_SEC, justify="left", anchor="w")
        self._stats_lbl.pack(padx=8, pady=(0, 6), anchor="w")

        # ── Check button ──────────────────────────────────────────────────
        self._chk_btn = ctk.CTkButton(
            body, text="開始查核", height=36, font=FT_BOLD,
            fg_color=COL_PROCESS, hover_color=COL_PROCESS_HOVER,
            text_color="#FFFFFF", command=self._verify)
        self._chk_btn.pack(fill="x", pady=(0, 2))

        self._chk_prog = ctk.CTkProgressBar(
            body, mode="indeterminate", height=4,
            fg_color=BG_INPUT, progress_color=COL_PROCESS)
        # 預設隱藏，查核時才顯示

        self._chk_lbl = ctk.CTkLabel(body, text="等待查核", font=FT_SM,
                                     text_color=TEXT_SEC)
        self._chk_lbl.pack()

        # ── Mark + Fix buttons ─────────────────────────────────────────────
        chk_btns = ctk.CTkFrame(body, fg_color="transparent")
        chk_btns.pack(fill="x", pady=(4, 4))
        # 標記問題＝輔助檢視 → ghost；一鍵修正＝查核流程收尾 → 處理藍（規則 1、2）
        self._mark_chk_btn = ctk.CTkButton(
            chk_btns, text="標記問題  →", height=28,
            font=FT_SM, fg_color=BG_INPUT, hover_color=BORDER,
            text_color=TEXT_PRI, text_color_disabled=TEXT_DIM,
            border_width=1, border_color=BORDER,
            state="disabled", command=self._mark_issues_in_final)
        self._mark_chk_btn.pack(side="left", fill="x", expand=True, padx=(0, 2))
        self._fix_chk_btn = ctk.CTkButton(
            chk_btns, text="一鍵修正", height=28,
            font=FT_SM, fg_color=COL_PROCESS, hover_color=COL_PROCESS_HOVER,
            text_color="#FFFFFF", text_color_disabled=TEXT_DIM,
            state="disabled", command=self._auto_fix_issues)
        self._fix_chk_btn.pack(side="left", fill="x", expand=True, padx=(2, 0))

        # ── Check results ─────────────────────────────────────────────────
        self._chk_box = _dark_text(body, height=5, font=("PingFang TC", 10))
        self._chk_box.pack(fill="both", expand=True, pady=(0, 2))
        self._chk_box.tag_configure("ok",   foreground=OK_FG,
                                    font=("PingFang TC", 10, "bold"))
        self._chk_box.tag_configure("err",  foreground=ERR_FG,
                                    font=("PingFang TC", 10, "bold"))
        self._chk_box.tag_configure("warn", foreground=WARN_FG,
                                    font=("PingFang TC", 10, "bold"))
        self._chk_box.tag_configure("info", foreground=INFO_FG,
                                    font=("PingFang TC", 10))
        self._chk_box.tag_configure("sub",  foreground=TEXT_SEC,
                                    font=("PingFang TC", 9))
        _add_copy_menu(self._chk_box)

    # ══════════════════════════════════════════════════════════════════════════
    #  COLUMN 3 — 最終稿件（可編輯）
    # ══════════════════════════════════════════════════════════════════════════
    def _build_col_output(self, parent):
        self._channel_header(parent, "最終稿件", COL_OUTPUT)

        body = ctk.CTkFrame(parent, fg_color="transparent")
        body.pack(fill="both", expand=True, padx=6, pady=6)

        # ── Author row ────────────────────────────────────────────────────
        self._col3_info_frame = info = ctk.CTkFrame(body, fg_color=BG_PANEL,
                                                     corner_radius=6)
        info.pack(fill="x", pady=(0, 4))
        r_auth = ctk.CTkFrame(info, fg_color="transparent")
        r_auth.pack(fill="x", padx=8, pady=6)
        ctk.CTkLabel(r_auth, text="記者", font=FT_SM,
                     text_color=TEXT_SEC).pack(side="left")
        self._author_var = tk.StringVar()
        auth_e = ctk.CTkEntry(r_auth, textvariable=self._author_var,
                              font=FT_SM, width=120,
                              fg_color=BG_INPUT, border_color=BORDER,
                              text_color=TEXT_PRI)
        auth_e.pack(side="left", padx=(6, 6))
        _add_entry_menu(auth_e)
        ctk.CTkCheckBox(r_auth, text="記憶", font=FT_SM,
                        text_color=TEXT_SEC, variable=self._remember_author,
                        fg_color=BORDER, hover_color=BORDER_LT,
                        checkmark_color=TEXT_PRI,
                        command=self._save_prefs).pack(side="left")

        # ── 主稿 / 配稿 tab toggle (hidden until sidebar mode) ─────────────
        self._col3_tab_frame = ctk.CTkFrame(body, fg_color=BG_PANEL,
                                             corner_radius=6, height=36)
        self._col3_tab_frame.pack_propagate(False)
        # Not packed initially — _toggle_sidebar controls this
        self._col3_main_btn = ctk.CTkButton(
            self._col3_tab_frame, text="主稿", font=FT_SM,
            fg_color=COL_OUTPUT, hover_color=COL_OUTPUT_HOVER,
            text_color="#FFFFFF", corner_radius=4, height=26, width=60,
            command=lambda: self._col3_switch_tab("main"))
        self._col3_main_btn.pack(side="left", padx=(4, 2), pady=5)
        self._col3_side_btn = ctk.CTkButton(
            self._col3_tab_frame, text="配稿", font=FT_SM,
            fg_color=BG_PANEL, hover_color=BORDER,
            text_color=TEXT_SEC, corner_radius=4, height=26, width=60,
            command=lambda: self._col3_switch_tab("side"))
        self._col3_side_btn.pack(side="left", padx=(0, 2), pady=5)

        # ── Main article container ─────────────────────────────────────────
        self._main_article_frame = ctk.CTkFrame(body, fg_color="transparent")
        self._main_article_frame.pack(fill="both", expand=True)

        r_title = ctk.CTkFrame(self._main_article_frame, fg_color="transparent")
        r_title.pack(fill="x", pady=(0, 2))
        ctk.CTkLabel(r_title, text="標題", font=FT_SM,
                     text_color=TEXT_SEC).pack(side="left")
        self._title_var = tk.StringVar()
        title_e = ctk.CTkEntry(r_title, textvariable=self._title_var,
                               font=FT, fg_color=BG_INPUT,
                               border_color=COL_OUTPUT, text_color=TEXT_PRI)
        title_e.pack(side="left", fill="x", expand=True, padx=(6, 0))
        _add_entry_menu(title_e)

        self._out_box = _dark_text(self._main_article_frame, height=8,
                                   font=FT_ARTICLE, undo=True,
                                   spacing1=2, spacing3=2)
        self._out_box.pack(fill="both", expand=True, pady=(0, 2))
        # 查核標記 tag：可修正（黃底）、僅供參考（灰底虛線）
        self._out_box.tag_configure("hl_fixable",
                                    background="#5C4A1E", foreground="#FFD966")
        self._out_box.tag_configure("hl_unsourced",
                                    background="#2A3040", foreground="#8899AA",
                                    underline=True)
        # 查核結果點擊跳轉時的暫時聚焦標記
        self._out_box.tag_configure("hl_focus",
                                    background="#2A4A6A", foreground="#FFFFFF")
        _add_copy_menu(self._out_box)
        self._out_box.bind("<<Modified>>", self._update_charcount)

        # ── Side article container (hidden by default) ─────────────────────
        self._side_article_frame = ctk.CTkFrame(body, fg_color="transparent")
        # Not packed initially

        r_side_title = ctk.CTkFrame(self._side_article_frame,
                                     fg_color="transparent")
        r_side_title.pack(fill="x", pady=(0, 2))
        ctk.CTkLabel(r_side_title, text="標題", font=FT_SM,
                     text_color=TEXT_SEC).pack(side="left")
        self._side_title_var = tk.StringVar()
        side_title_e = ctk.CTkEntry(
            r_side_title, textvariable=self._side_title_var, font=FT,
            fg_color=BG_INPUT, border_color=COL_OUTPUT, text_color=TEXT_PRI)
        side_title_e.pack(side="left", fill="x", expand=True, padx=(6, 0))
        _add_entry_menu(side_title_e)

        self._side_box = _dark_text(self._side_article_frame, height=8,
                                    font=FT_ARTICLE, undo=True,
                                    spacing1=2, spacing3=2)
        self._side_box.pack(fill="both", expand=True, pady=(0, 2))
        _add_copy_menu(self._side_box)
        self._side_box.bind("<<Modified>>", self._update_charcount)

        # ── 字數（always visible, updates per active tab）─────────────────
        self._charcount_lbl = ctk.CTkLabel(
            body, text="0 字", font=FT_SM, text_color=TEXT_SEC, anchor="e")
        self._charcount_lbl.pack(fill="x", padx=4, pady=(0, 2))

        # ── Export buttons (bottom) ───────────────────────────────────────
        self._export_bar = eb = ctk.CTkFrame(body, fg_color="transparent")
        eb.pack(fill="x", pady=(4, 0))
        for label, cmd in [("複製", self._copy),
                           ("TXT", self._save_txt),
                           ("DOCX", self._save_docx)]:
            ctk.CTkButton(eb, text=label, width=52, height=26,
                          font=FT_SM, fg_color=COL_OUTPUT,
                          hover_color=COL_OUTPUT_HOVER, text_color="#FFFFFF",
                          command=cmd).pack(side="left", padx=(0, 3))
        ctk.CTkButton(eb, text="清除", width=44, height=26,
                      font=FT_SM, fg_color=BG_INPUT, hover_color=BORDER,
                      text_color=TEXT_SEC, border_width=1, border_color=BORDER,
                      command=self._clear_output).pack(side="left", padx=(3, 0))

    # ══════════════════════════════════════════════════════════════════════════
    #  WAVE ANIMATION
    # ══════════════════════════════════════════════════════════════════════════
    def _wave_tick(self):
        if not self._wave_active:
            self._wave_canvas.delete("all")
            self._wave_anim_id = None
            return
        w = self._wave_canvas.winfo_width()
        h = self._wave_canvas.winfo_height()
        if w < 4:
            self._wave_anim_id = self.root.after(50, self._wave_tick)
            return
        self._wave_frame += 2
        step = 3
        pts  = []
        for x in range(0, w + step, step):
            phase = (x / 38.0 + self._wave_frame * 0.14) * math.pi
            amp   = (h / 2 - 3) * 0.75
            y     = h / 2 + amp * math.sin(phase)
            pts.extend([x, y])
        self._wave_canvas.delete("all")
        if len(pts) >= 4:
            self._wave_canvas.create_line(*pts, fill=ACCENT_BLUE,
                                          smooth=True, width=2)
        self._wave_anim_id = self.root.after(45, self._wave_tick)

    def _wave_start(self):
        self._wave_active = True
        if not self._wave_anim_id:
            self._wave_tick()

    def _wave_stop(self):
        self._wave_active = False

    # ── record button pulse ───────────────────────────────────────────────
    def _rec_pulse(self):
        if not self._wave_active:
            self._start_btn.itemconfig(self._rec_circle, fill=ACCENT_RED)
            self._rec_anim_id = None
            return
        self._rec_pulse_on = not self._rec_pulse_on
        color = "#FF5252" if self._rec_pulse_on else "#B71C1C"
        self._start_btn.itemconfig(self._rec_circle, fill=color)
        self._rec_anim_id = self.root.after(500, self._rec_pulse)

    def _rec_pulse_start(self):
        if not self._rec_anim_id:
            self._rec_pulse()

    # ══════════════════════════════════════════════════════════════════════════
    #  COLUMN 1 HANDLERS
    # ══════════════════════════════════════════════════════════════════════════
    def _set_phone_mode(self, on: bool):
        """進入／離開手機麥克風分頁時啟停遠端麥克風伺服器。"""
        if on:
            self._phone_mode.set(True)
            ip  = _get_local_ip()
            url = f"https://{ip}:8765"
            self._phone_url_var.set(url)
            self._update_qr(url)
            ensure_server_running()
            self._arm_phone_mic()
        else:
            self._phone_mode.set(False)
            if self._transcriber:
                self._transcriber.stop()
                self._transcriber = None
            self._wave_stop()
            self._live_lbl.configure(text="", text_color=TEXT_SEC)

    def _arm_phone_mic(self):
        if not HAS_PHONE_MIC or not HAS_MLX_WHISPER:
            return
        if self._transcriber:
            return

        def _st(s):  self.root.after(0, lambda: self._live_lbl.configure(
                         text=s, text_color=TEXT_SEC))
        def _err(e): self.root.after(0, lambda: self._live_lbl.configure(
                         text=f"錯誤：{e}", text_color=ERR_FG))
        def _done():    self.root.after(0, self._live_done)
        def _started(): self.root.after(0, self._phone_did_start)
        def _stopped(): self.root.after(0, self._phone_did_stop)

        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        rec_dir = os.path.join(SCRIPT_DIR, "rec")
        os.makedirs(rec_dir, exist_ok=True)
        self._transcriber = PhoneMicTranscriber(
            model_name=self._wmodel_var.get(),
            language=self._lang_var.get(),
            on_segment=self._on_segment, on_status=_st,
            on_error=_err, on_done=_done,
            on_phone_started=_started,
            on_phone_stopped=_stopped,
            audio_save_path=os.path.join(rec_dir, f"phone_{ts}.wav"))
        self._transcriber.start()
        self._live_lbl.configure(text="模型載入中...", text_color=TEXT_SEC)
        self._status("手機麥克風模式：備妥中")

    def _phone_did_stop(self):
        if self._wave_active:
            self._live_stop()

    def _phone_did_start(self):
        if self._wave_active:
            return
        # 新段開始前清除上一次的逐字稿（上一次已由 _autosave_transcript 自動存檔）
        self._transcript.delete("1.0", tk.END)
        log_dir = os.path.join(SCRIPT_DIR, "log")
        os.makedirs(log_dir, exist_ok=True)
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        self._log_file = open(
            os.path.join(log_dir, f"phone_{ts}.txt"), "w", encoding="utf-8")
        self._wave_start()
        self._rec_pulse_start()
        self._stop_btn.configure(cursor="hand2")
        self._stop_btn.itemconfig(self._stop_square, fill=ACCENT_RED)
        self._live_lbl.configure(text="錄音中（手機）...", text_color=ERR_FG)
        self._status("手機麥克風轉錄中")

    def _update_qr(self, url: str):
        photo = make_qr_image_tk(url)
        if photo:
            self._qr_photo = photo
            self._qr_label.configure(image=self._qr_photo, text="")
        else:
            self._qr_label.configure(
                image="",
                text="（安裝 qrcode + Pillow 後可顯示 QR Code）",
                font=FT_SM, fg=TEXT_SEC, bg=BG_PANEL)

    def _copy_phone_url(self):
        url = self._phone_url_var.get()
        if url:
            self.root.clipboard_clear()
            self.root.clipboard_append(url)
            self._status("已複製手機網址")

    def _toggle_sidebar(self):
        if self._sidebar_mode.get():
            self._sidebar_ctx_frame.pack(fill="x",
                                         after=self._sidebar_chk_row)
            # 顯示欄3 tab 選擇器，並切換到主稿
            self._col3_tab_frame.pack(fill="x", pady=(0, 4),
                                       after=self._col3_info_frame)
            self._col3_switch_tab("main")
        else:
            self._sidebar_ctx_frame.pack_forget()
            # 隱藏 tab 選擇器，回到單一主稿視圖
            self._col3_tab_frame.pack_forget()
            self._side_article_frame.pack_forget()
            if not self._main_article_frame.winfo_ismapped():
                self._main_article_frame.pack(fill="both", expand=True,
                                               before=self._charcount_lbl)
            self._col3_active = "main"

    def _col3_switch_tab(self, tab_name: str):
        self._col3_active = tab_name
        if tab_name == "main":
            self._side_article_frame.pack_forget()
            if not self._main_article_frame.winfo_ismapped():
                self._main_article_frame.pack(fill="both", expand=True,
                                               before=self._charcount_lbl)
            self._col3_main_btn.configure(fg_color=COL_OUTPUT,
                                           text_color="#FFFFFF")
            self._col3_side_btn.configure(fg_color=BG_PANEL,
                                           text_color=TEXT_SEC)
        else:
            self._main_article_frame.pack_forget()
            if not self._side_article_frame.winfo_ismapped():
                self._side_article_frame.pack(fill="both", expand=True,
                                               before=self._charcount_lbl)
            self._col3_main_btn.configure(fg_color=BG_PANEL,
                                           text_color=TEXT_SEC)
            self._col3_side_btn.configure(fg_color=COL_OUTPUT,
                                           text_color="#FFFFFF")
        self._update_charcount()

    def _refresh_devices(self):
        if not HAS_SOUNDDEVICE:
            self._dev_combo.configure(values=["（請安裝 sounddevice）"])
            self._dev_combo.set("（請安裝 sounddevice）")
            return
        import sounddevice as sd
        inputs = [(i, d["name"]) for i, d in enumerate(sd.query_devices())
                  if d["max_input_channels"] > 0]
        self._device_map = {n: i for i, n in inputs}
        names = [n for _, n in inputs]
        self._dev_combo.configure(values=names)
        if names:
            self._dev_combo.set(names[0])

    def _live_start(self):
        if self._wave_active: return
        if self._phone_mode.get():
            return
        if not HAS_SOUNDDEVICE or not HAS_MLX_WHISPER:
            messagebox.showerror("缺少套件",
                                  "請執行：pip install sounddevice mlx-whisper")
            return
        if not HAS_LIVE:
            messagebox.showerror("錯誤",
                                  "無法載入 LiveTranscriber\n請確認 gui_whisper_live.py 在同目錄")
            return
        dev = self._device_map.get(self._dev_var.get())
        def _st(s):  self.root.after(0, lambda: self._live_lbl.configure(
                         text=s, text_color=COL_INPUT))
        def _err(e): self.root.after(0, lambda: self._live_lbl.configure(
                         text=f"錯誤：{e}", text_color=ERR_FG))
        def _done(): self.root.after(0, self._live_done)
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        log_dir = os.path.join(SCRIPT_DIR, "log")
        os.makedirs(log_dir, exist_ok=True)
        rec_dir = os.path.join(SCRIPT_DIR, "rec")
        os.makedirs(rec_dir, exist_ok=True)
        self._transcriber = LiveTranscriber(
            device_index=dev, model_name=self._wmodel_var.get(),
            language=self._lang_var.get(), diarize=False, hf_token="",
            on_segment=self._on_segment, on_status=_st, on_error=_err,
            on_done=_done,
            audio_save_path=os.path.join(rec_dir, f"live_{ts}.wav"))
        self._log_file = open(
            os.path.join(log_dir, f"live_{ts}.txt"), "w", encoding="utf-8")
        self._transcriber.start()
        self._wave_start()
        self._rec_pulse_start()
        self._stop_btn.configure(cursor="hand2")
        self._stop_btn.itemconfig(self._stop_square, fill=ACCENT_RED)
        self._live_lbl.configure(text="錄音中...", text_color=ERR_FG)
        self._status("即時轉錄中")

    def _live_stop(self):
        if self._transcriber: self._transcriber.stop()
        self._live_lbl.configure(text="停止中...", text_color=TEXT_SEC)

    def _live_done(self):
        if self._log_file:
            self._log_file.close()
            self._log_file = None
        # 在 transcriber 清空前先抓 scene 與 WAV 路徑（手機模式才有）
        scene = ""
        wav_path = None
        if self._transcriber and hasattr(self._transcriber, "last_scene"):
            scene = self._transcriber.last_scene
        if self._transcriber and hasattr(self._transcriber, "_audio_save_path"):
            wav_path = self._transcriber._audio_save_path
        # 若有 scene，把 WAV 檔名加上記者會主題
        if wav_path and scene and os.path.exists(wav_path):
            safe = re.sub(r'[\\/*?:"<>|\s]+', "_", scene)[:32].strip("_")
            new_wav = wav_path.replace(".wav", f"_{safe}.wav")
            try:
                os.rename(wav_path, new_wav)
            except Exception:
                pass
        self._wave_stop()
        self._stop_btn.configure(cursor="arrow")
        self._stop_btn.itemconfig(self._stop_square, fill=TEXT_DIM)
        self._start_btn.itemconfig(self._rec_circle, fill=ACCENT_RED)
        if scene:
            self._last_scene = scene   # 保留供校正階段使用
        self._transcriber = None
        self._live_lbl.configure(text="已停止", text_color=TEXT_SEC)
        # 自動儲存逐字稿完整快照（含手動貼上或編輯的內容）
        saved = self._autosave_transcript(scene=scene)
        self._status(f"就緒｜已自動儲存：{saved}" if saved else "就緒")
        if self._phone_mode.get() and HAS_PHONE_MIC:
            self.root.after(300, self._arm_phone_mic)

    def _autosave_transcript(self, scene: str = "") -> str:
        """儲存逐字稿框完整內容，回傳檔名；內容為空時回傳空字串。"""
        text = self._transcript.get("1.0", tk.END).strip()
        if not text:
            return ""
        try:
            log_dir = os.path.join(SCRIPT_DIR, "log")
            os.makedirs(log_dir, exist_ok=True)
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            safe = re.sub(r'[\\/*?:"<>|\s]+', "_", scene)[:24].strip("_")
            name = f"transcript_{ts}_{safe}.txt" if safe else f"transcript_{ts}.txt"
            with open(os.path.join(log_dir, name), "w", encoding="utf-8") as f:
                f.write(text)
            return f"log/{name}"
        except Exception:
            return ""

    def _live_clear(self):
        self._transcript.delete("1.0", tk.END)

    def _on_segment(self, text: str, speaker):
        def _ins():
            ts = datetime.datetime.now().strftime("%H:%M:%S")
            self._transcript.insert(tk.END, f"[{ts}]  {text}\n", "body")
            self._transcript.see(tk.END)
            if self._log_file:
                self._log_file.write(f"[{ts}]  {text}\n")
                self._log_file.flush()
        self.root.after(0, _ins)

    def _paste_transcript(self):
        try:
            text = self.root.clipboard_get()
            if text.strip():
                self._transcript.insert(tk.END, text)
                self._transcript.see(tk.END)
                self._status("已貼上逐字稿")
        except tk.TclError:
            pass

    def _load_transcript(self):
        path = filedialog.askopenfilename(
            title="載入逐字稿",
            filetypes=[("文字檔", "*.txt"), ("所有", "*.*")])
        if not path: return
        try:
            with open(path, encoding="utf-8", errors="replace") as f:
                text = f.read()
            self._transcript.delete("1.0", tk.END)
            self._transcript.insert(tk.END, text)
            self._transcript.see(tk.END)
            self._status(f"已載入：{os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("載入失敗", str(e))

    def _save_transcript(self):
        text = self._transcript.get("1.0", tk.END).strip()
        if not text:
            messagebox.showwarning("空白", "逐字稿內容為空"); return
        ts   = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        path = filedialog.asksaveasfilename(
            defaultextension=".txt", initialfile=f"transcript_{ts}.txt",
            filetypes=[("文字檔", "*.txt"), ("所有", "*.*")])
        if not path: return
        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(text)
            self._status(f"已儲存：{os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("儲存失敗", str(e))

    def _correct_transcript(self):
        if self._wave_active:
            messagebox.showwarning("錄音中", "請先停止錄音再校正逐字稿")
            return
        text = self._transcript.get("1.0", tk.END).strip()
        if not text:
            messagebox.showwarning("空白", "逐字稿內容為空，無法校正")
            return
        api_key = self._apikey_var.get().strip()
        if not api_key:
            api_key = os.environ.get("ANTHROPIC_API_KEY", "")
        if not api_key:
            messagebox.showwarning("缺少 API Key",
                                   "請在右上角「⚙ 設定」填入 Anthropic API Key")
            return

        # 組合背景資訊：受訪者 + 最近一次記者會名稱
        context_parts = []
        interviewee = self._interviewee_var.get().strip()
        if interviewee:
            context_parts.append(f"受訪者 / 發言人：{interviewee}")
        if self._last_scene:
            context_parts.append(f"記者會名稱：{self._last_scene}")
        context = "\n".join(context_parts)

        self._correct_btn.configure(state="disabled", text="修正中...")
        self._status("逐字稿錯字修正中...")

        def _done(corrected):
            def _ui():
                self._transcript.delete("1.0", tk.END)
                self._transcript.insert(tk.END, corrected)
                self._correct_btn.configure(state="normal", text="修正錯字")
                self._status("逐字稿錯字修正完成")
            self.root.after(0, _ui)

        def _err(msg):
            def _ui():
                self._correct_btn.configure(state="normal", text="修正錯字")
                self._status(f"修正失敗：{msg[:60]}")
                messagebox.showerror("修正失敗", msg)
            self.root.after(0, _ui)

        claude_correct(text, api_key, _done, _err, context=context)

    def _translate_transcript(self):
        if self._wave_active:
            messagebox.showwarning("錄音中", "請先停止錄音再翻譯逐字稿")
            return
        text = self._transcript.get("1.0", tk.END).strip()
        if not text:
            messagebox.showwarning("空白", "逐字稿內容為空，無法翻譯")
            return
        api_key = self._apikey_var.get().strip()
        if not api_key:
            api_key = os.environ.get("ANTHROPIC_API_KEY", "")
        if not api_key:
            messagebox.showwarning("缺少 API Key",
                                   "請在右上角「⚙ 設定」填入 Anthropic API Key")
            return

        self._trans_btn.configure(state="disabled", text="翻譯中...")
        self._status("逐字稿翻譯中...")

        def _done(translated):
            def _ui():
                self._transcript.delete("1.0", tk.END)
                self._transcript.insert(tk.END, translated)
                self._trans_btn.configure(state="normal", text="翻成中文")
                self._status("逐字稿已翻譯為繁體中文")
            self.root.after(0, _ui)

        def _err(msg):
            def _ui():
                self._trans_btn.configure(state="normal", text="翻成中文")
                self._status(f"翻譯失敗：{msg[:60]}")
                messagebox.showerror("翻譯失敗", msg)
            self.root.after(0, _ui)

        claude_translate(text, api_key, _done, _err)

    def _audio_browse(self):
        path = filedialog.askopenfilename(
            title="選擇音檔",
            filetypes=[
                ("音訊檔案", "*.mp3 *.wav *.m4a *.ogg *.flac *.aac *.wma "
                             "*.mp4 *.mkv *.mov *.webm"),
                ("MP3", "*.mp3"), ("WAV", "*.wav"),
                ("M4A", "*.m4a"), ("FLAC", "*.flac"), ("所有", "*.*"),
            ])
        if path:
            self._audio_file = path
            self._audio_name_var.set(os.path.basename(path))
            self._audio_lbl.configure(text="", text_color=TEXT_SEC)

    def _audio_transcribe(self):
        if self._audio_busy: return
        if not self._audio_file:
            messagebox.showwarning("缺少音檔", "請先選擇音檔"); return
        if not HAS_MLX_WHISPER:
            messagebox.showerror("缺少套件",
                                  "請執行：pip install mlx-whisper"); return

        self._audio_busy = True
        self._audio_btn.configure(state="disabled", text="轉錄中...")
        self._audio_prog.pack(side="left", padx=(8, 4), after=self._audio_btn)
        self._audio_prog.start()
        self._audio_lbl.configure(text="載入模型...", text_color=COL_INPUT)
        self._status("音檔轉錄中")

        path       = self._audio_file
        model_name = self._wmodel_var.get()
        language   = self._lang_var.get()
        if language == "auto": language = None

        def _set_lbl(msg, color=COL_INPUT):
            self.root.after(0, lambda: self._audio_lbl.configure(
                text=msg, text_color=color))

        def _run():
            try:
                import mlx_whisper
                from gui_whisper_live import MLX_MODEL_MAP
                repo = MLX_MODEL_MAP.get(model_name, f"mlx-community/whisper-{model_name}")
                _set_lbl("載入模型...")
                result = mlx_whisper.transcribe(
                    path,
                    path_or_hf_repo=repo,
                    language=language,
                    verbose=False,
                )
                segments = result.get("segments", [])
                fname = os.path.basename(path)
                self.root.after(0, lambda: (
                    self._transcript.insert(
                        tk.END, f"\n--- 音檔：{fname} ---\n", "body"),
                    self._transcript.see(tk.END)
                ))
                for seg in segments:
                    text = seg["text"].strip()
                    if not text: continue
                    m, s  = int(seg["start"]) // 60, int(seg["start"]) % 60
                    stamp = f"{m:02d}:{s:02d}"
                    def _ins(t=text, ts=stamp):
                        self._transcript.insert(
                            tk.END, f"[{ts}]  {t}\n", "body")
                        self._transcript.see(tk.END)
                        if self._log_file:
                            self._log_file.write(f"[{ts}]  {t}\n")
                            self._log_file.flush()
                    self.root.after(0, _ins)

                try:
                    log_dir  = os.path.join(SCRIPT_DIR, "log")
                    os.makedirs(log_dir, exist_ok=True)
                    ts       = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                    base     = os.path.splitext(fname)[0][:40]
                    log_path = os.path.join(log_dir, f"audio_{base}_{ts}.txt")
                    content  = self._transcript.get("1.0", tk.END)
                    with open(log_path, "w", encoding="utf-8") as lf:
                        lf.write(content)
                    saved_name = os.path.basename(log_path)
                except Exception:
                    saved_name = ""

                def _ok(sn=saved_name):
                    self._audio_busy = False
                    self._audio_btn.configure(state="normal", text="開始轉錄")
                    self._audio_prog.stop()
                    self._audio_prog.pack_forget()
                    hint = f"完成，已備份 {sn}" if sn else "轉錄完成"
                    self._audio_lbl.configure(text=hint, text_color=OK_FG)
                    self._status(f"音檔轉錄完成 | 備份：log/{sn}")
                self.root.after(0, _ok)
            except Exception as e:
                def _fail(msg=str(e)):
                    self._audio_busy = False
                    self._audio_btn.configure(state="normal", text="開始轉錄")
                    self._audio_prog.stop()
                    self._audio_prog.pack_forget()
                    self._audio_lbl.configure(
                        text=f"錯誤：{msg[:40]}", text_color=ERR_FG)
                    self._status(f"轉錄失敗：{msg[:60]}")
                self.root.after(0, _fail)

        threading.Thread(target=_run, daemon=True).start()

    def _file_browse(self):
        paths = filedialog.askopenfilenames(
            title="選擇附件",
            filetypes=[
                ("支援格式", "*.pdf *.docx *.doc *.xlsx *.xls *.txt"),
                ("PDF", "*.pdf"), ("Word", "*.docx *.doc"),
                ("Excel", "*.xlsx *.xls"), ("文字", "*.txt"), ("所有", "*.*"),
            ])
        for p in paths:
            fname = os.path.basename(p)
            if any(f == fname for f, _ in self._files): continue
            self._status(f"解析：{fname}...")
            self._files.append((fname, parse_file(p)))
            self._flist.insert(tk.END, f"  {fname}")
        self._status(f"已載入 {len(self._files)} 個附件")

    def _file_preview(self, _=None):
        sel = self._flist.curselection()
        if not sel: return
        fname, text = self._files[sel[0]]
        w = ctk.CTkToplevel(self.root)
        w.title(f"預覽：{fname}")
        w.geometry("700x500")
        st = _dark_text(w, font=FT_MONO)
        st.pack(fill="both", expand=True, padx=6, pady=6)
        st.insert(tk.END, text)
        _add_copy_menu(st)

    def _file_remove(self):
        sel = self._flist.curselection()
        if not sel: return
        self._flist.delete(sel[0]); del self._files[sel[0]]

    def _file_clear(self):
        self._flist.delete(0, tk.END); self._files.clear()

    # ══════════════════════════════════════════════════════════════════════════
    #  COLUMN 2 HANDLERS
    # ══════════════════════════════════════════════════════════════════════════
    def _update_key_label(self):
        has = bool(self._apikey_var.get().strip())
        self._topbar_key_lbl.configure(
            text="API Key: 已設定" if has else "API Key: 未設定",
            text_color=OK_FG if has else ERR_FG)

    def _toggle_apikey_visibility(self):
        self._apikey_visible = not self._apikey_visible
        ent = self._apikey_entry
        if ent is not None and ent.winfo_exists():
            ent.configure(show="" if self._apikey_visible else "*")

    def _save_prefs(self):
        data = {}
        if self._remember_key.get():
            key = self._apikey_var.get().strip()
            # 防呆：偵測並移除重複貼上的 key
            half = len(key) // 2
            if half > 20 and key[:half] == key[half:]:
                key = key[:half]
                self._apikey_var.set(key)
            data["api_key"] = key
        else:
            data["api_key"] = ""
        if self._remember_author.get():
            data["author"] = self._author_var.get().strip()
        else:
            data["author"] = ""
        save_config(data)

    # ══════════════════════════════════════════════════════════════════════════
    #  SETTINGS DIALOG（API Key / 模型 / 技能包 / Whisper）
    # ══════════════════════════════════════════════════════════════════════════
    def _open_settings(self):
        win = self._settings_win
        if win is not None and win.winfo_exists():
            win.lift()
            win.focus_force()
            return

        win = ctk.CTkToplevel(self.root)
        self._settings_win = win
        win.title("設定")
        win.configure(fg_color=BG)
        win.transient(self.root)
        win.resizable(False, False)
        x = self.root.winfo_rootx() + 100
        y = self.root.winfo_rooty() + 80
        win.geometry(f"560x440+{x}+{y}")

        def _section(text):
            ctk.CTkLabel(win, text=text, font=FT_BOLD,
                         text_color=TEXT_PRI).pack(anchor="w",
                                                   padx=16, pady=(16, 2))
            ctk.CTkFrame(win, fg_color=BORDER, height=1
                         ).pack(fill="x", padx=16, pady=(0, 4))

        def _row(label):
            r = ctk.CTkFrame(win, fg_color="transparent")
            r.pack(fill="x", padx=16, pady=3)
            ctk.CTkLabel(r, text=label, font=FT_SM, text_color=TEXT_SEC,
                         width=60, anchor="w").pack(side="left")
            return r

        # ── Claude 寫稿 ────────────────────────────────────────────────────
        _section("Claude 寫稿")

        r = _row("API Key")
        self._apikey_entry = ctk.CTkEntry(
            r, textvariable=self._apikey_var,
            show="" if self._apikey_visible else "*",
            font=("Menlo", 9), fg_color=BG_INPUT, border_color=BORDER,
            text_color=TEXT_PRI)
        self._apikey_entry.pack(side="left", fill="x", expand=True, padx=(4, 4))
        _add_entry_menu(self._apikey_entry)
        self._ghost_btn(r, "顯示", self._toggle_apikey_visibility,
                        width=36, height=24).pack(side="left")

        r = _row("")
        ctk.CTkCheckBox(r, text="記憶 API Key（儲存於 config.json）",
                        font=FT_SM, text_color=TEXT_SEC,
                        variable=self._remember_key,
                        fg_color=BORDER, hover_color=BORDER_LT,
                        checkmark_color=TEXT_PRI,
                        command=self._save_prefs).pack(side="left", padx=4)

        r = _row("模型")
        self._make_combo(r, self._gen_model,
                         ["claude-sonnet-4-6", "claude-opus-4-6",
                          "claude-haiku-4-5-20251001"],
                         width=240).pack(side="left", padx=(4, 0))

        r = _row("技能包")
        self._skill_lbl = ctk.CTkLabel(r, text="", font=FT_SM,
                                       text_color=OK_FG, anchor="w")
        self._skill_lbl.pack(side="left", fill="x", expand=True, padx=4)
        self._ghost_btn(r, "選擇", self._browse_skill,
                        width=36, height=24).pack(side="left")
        self._update_skill_label()

        # ── 語音轉錄 ───────────────────────────────────────────────────────
        _section("語音轉錄（Whisper）")

        r = _row("模型")
        self._make_combo(r, self._wmodel_var, WHISPER_MODELS,
                         width=160).pack(side="left", padx=(4, 16))
        ctk.CTkLabel(r, text="語言", font=FT_SM,
                     text_color=TEXT_SEC).pack(side="left")
        self._make_combo(r, self._lang_var, LANGUAGES,
                         width=80).pack(side="left", padx=(4, 0))

        ctk.CTkLabel(win, text="現場錄音、手機麥克風與音檔轉錄共用此設定；"
                               "首次選用新模型會自動下載（需網路）。",
                     font=FT_SM, text_color=TEXT_DIM, anchor="w"
                     ).pack(fill="x", padx=16, pady=(2, 0))

        # ── 關閉 ───────────────────────────────────────────────────────────
        def _close():
            self._save_prefs()
            self._skill_lbl    = None
            self._apikey_entry = None
            self._settings_win = None
            win.destroy()

        ctk.CTkButton(win, text="完成", height=34, font=FT_BOLD,
                      fg_color=ACCENT_BLUE, hover_color=COL_PROCESS_HOVER,
                      text_color="#FFFFFF", command=_close
                      ).pack(side="bottom", fill="x", padx=16, pady=14)
        win.protocol("WM_DELETE_WINDOW", _close)
        win.after(80, win.lift)

    def _generate(self):
        if self._generating: return
        self._generating = True
        sidebar = self._sidebar_mode.get()
        label   = "生成主稿＋配稿中..." if sidebar else "生成中..."
        self._gen_btn.configure(state="disabled", text=label)
        self._gen_prog.pack(fill="x", pady=(0, 2), before=self._gen_lbl)
        self._gen_prog.start()
        self._gen_lbl.configure(text="呼叫 Claude API...", text_color=COL_PROCESS)

        # Clear output boxes before streaming
        self._out_box.delete("1.0", tk.END)
        self._title_var.set("")
        if sidebar:
            self._side_box.delete("1.0", tk.END)
            self._side_title_var.set("")

        transcript = self._transcript.get("1.0", tk.END).strip()
        files_text = "\n\n".join(f"=== {n} ===\n{t}" for n, t in self._files)

        def _tok_main(c):
            self.root.after(0, lambda t=c: (
                self._out_box.insert(tk.END, t),
                self._out_box.see(tk.END)
            ))

        def _tok_side(c):
            self.root.after(0, lambda t=c: (
                self._side_box.insert(tk.END, t),
                self._side_box.see(tk.END)
            ))

        def _done():
            def _ui():
                self._generating = False
                self._gen_btn.configure(state="normal", text="生成報導")
                self._gen_prog.stop()
                self._gen_prog.pack_forget()

                # 只有確實有內容時才宣告成功，否則保留錯誤訊息
                raw_main = self._out_box.get("1.0", tk.END).strip()
                if not raw_main:
                    return

                self._gen_lbl.configure(text="生成完成", text_color=OK_FG)
                self._status("報導生成完成")

                # Extract title from main box and clean up
                t, b = parse_single_output(strip_markdown(raw_main))
                self._title_var.set(t)
                self._out_box.delete("1.0", tk.END)
                if b:
                    self._out_box.insert(tk.END, b)

                # Extract title from side box if in sidebar mode
                if sidebar:
                    raw_side = self._side_box.get("1.0", tk.END).strip()
                    st, sb = parse_single_output(strip_markdown(raw_side))
                    self._side_title_var.set(st)
                    self._side_box.delete("1.0", tk.END)
                    if sb:
                        self._side_box.insert(tk.END, sb)

                self._out_box.edit_modified(True)
                self._update_charcount()
            self.root.after(0, _ui)

        def _err(msg):
            def _ui():
                self._generating = False
                self._gen_btn.configure(state="normal", text="生成報導")
                self._gen_prog.stop()
                self._gen_prog.pack_forget()
                self._gen_lbl.configure(
                    text=f"錯誤：{msg[:40]}", text_color=ERR_FG)
                self._status(f"生成失敗：{msg[:60]}")
            self.root.after(0, _ui)

        claude_generate(
            transcript, files_text,
            self._gen_model.get(),
            self._interviewee_var.get().strip(),
            self._main_angle_var.get().strip(),
            sidebar,
            self._side_angle_var.get().strip(),
            self._apikey_var.get().strip(),
            self._skill_path_var.get(),
            _tok_main, _tok_side, _done, _err
        )

    # ══════════════════════════════════════════════════════════════════════════
    #  FACT CHECK HANDLERS
    # ══════════════════════════════════════════════════════════════════════════
    def _verify(self):
        if self._checking: return
        article = self._out_box.get("1.0", tk.END).strip()
        if not article:
            messagebox.showwarning("缺少稿件",
                "最終稿件欄為空，請先生成報導，再進行查核")
            return
        source = ("\n\n".join(f"=== {n} ===\n{t}" for n, t in self._files)
                  or self._transcript.get("1.0", tk.END).strip())

        self._out_box.tag_remove("hl_fixable", "1.0", tk.END)
        self._out_box.tag_remove("hl_unsourced", "1.0", tk.END)
        self._chk_box.delete("1.0", tk.END)
        self._chk_box.insert(tk.END, "數字出處比對\n", "info")

        art_nums   = extract_numbers(article)
        src_nums   = extract_numbers(source)
        verified   = art_nums & src_nums
        unverified = art_nums - src_nums

        if verified:
            self._chk_box.insert(tk.END,
                f"[OK]  有出處：{len(verified)} 個數字\n", "ok")
        if unverified:
            self._chk_box.insert(tk.END,
                f"[!]   找不到出處（需人工確認）：\n", "warn")
            for n in sorted(unverified)[:12]:
                self._chk_box.insert(tk.END, f"      {n}\n", "warn")

        self._stats_lbl.configure(
            text=(f"稿件數字：{len(art_nums)} 個\n"
                  f"  有出處：{len(verified)} 個\n"
                  f"  待確認：{len(unverified)} 個"))

        if not source.strip(): return
        if not HAS_ANTHROPIC:
            self._chk_box.insert(tk.END,
                "\n（安裝 anthropic SDK 後可啟用語義查核）\n", "warn")
            return

        self._checking = True
        self._chk_btn.configure(state="disabled", text="查核中...")
        self._chk_prog.pack(fill="x", pady=(0, 2), before=self._chk_lbl)
        self._chk_prog.start()
        self._chk_lbl.configure(text="Haiku 語義查核中...",
                                text_color=COL_PROCESS)

        def _res(issues: list):
            def _ui():
                self._last_check_issues = issues
                self._chk_box.insert(tk.END, "\nClaude 查核結果\n", "info")
                if not issues:
                    self._chk_box.insert(tk.END,
                        "[OK]  稿件內容與素材吻合，未發現問題\n", "ok")
                    self._mark_chk_btn.configure(state="disabled")
                    self._fix_chk_btn.configure(state="disabled")
                else:
                    self._chk_box.insert(tk.END,
                        "（點擊問題文字可跳至稿件對應位置）\n", "sub")
                    fixable = 0
                    for idx, item in enumerate(issues):
                        val  = item.get("value", "")
                        sug  = item.get("suggestion", "")
                        iss  = item.get("issue", "")
                        typ  = item.get("type", "?")
                        if typ == "typo":
                            tag, prefix = "warn", "錯字"
                        elif typ == "mismatch":
                            tag, prefix = "err", "不符"
                        else:
                            tag, prefix = "info", "無出處"
                        jump_tag = f"jump_{idx}"
                        self._chk_box.insert(tk.END,
                            f"[{prefix}]  「{val}」\n", (tag, jump_tag))
                        self._chk_box.tag_bind(
                            jump_tag, "<Button-1>",
                            lambda e, v=val: self._jump_to_issue(v))
                        self._chk_box.tag_bind(
                            jump_tag, "<Enter>",
                            lambda e: self._chk_box.configure(cursor="hand2"))
                        self._chk_box.tag_bind(
                            jump_tag, "<Leave>",
                            lambda e: self._chk_box.configure(cursor="xterm"))
                        if sug:
                            self._chk_box.insert(tk.END,
                                f"  → 建議：{sug}\n", "ok")
                            fixable += 1
                        if iss:
                            self._chk_box.insert(tk.END,
                                f"  {iss}\n", "sub")
                    self._mark_chk_btn.configure(state="normal")
                    if fixable > 0:
                        self._fix_chk_btn.configure(
                            state="normal",
                            text=f"一鍵修正（{fixable} 項）")
                    else:
                        self._fix_chk_btn.configure(state="disabled")
            self.root.after(0, _ui)

        def _done():
            def _ui():
                self._checking = False
                self._chk_btn.configure(state="normal", text="開始查核")
                self._chk_prog.stop()
                self._chk_prog.pack_forget()
                self._chk_lbl.configure(text="查核完成", text_color=OK_FG)
                self._status("事實查核完成")
            self.root.after(0, _ui)

        def _err(msg):
            def _ui():
                self._checking = False
                self._chk_btn.configure(state="normal", text="開始查核")
                self._chk_prog.stop()
                self._chk_prog.pack_forget()
                self._chk_box.insert(tk.END, f"[X] {msg}\n", "err")
                self._chk_lbl.configure(text="查核失敗", text_color=ERR_FG)
            self.root.after(0, _ui)

        claude_check(article, source, self._apikey_var.get().strip(),
                     _res, _done, _err)

    def _jump_to_issue(self, val: str):
        """點擊查核結果中的問題文字，跳至最終稿件對應位置並短暫聚焦標記。"""
        if not val:
            return
        if self._col3_active != "main":
            self._col3_switch_tab("main")
        pos = self._out_box.search(val, "1.0", tk.END)
        if not pos:
            self._status("稿件中找不到這段文字（可能已被修改）")
            return
        end = f"{pos}+{len(val)}c"
        self._out_box.tag_remove("hl_focus", "1.0", tk.END)
        self._out_box.tag_add("hl_focus", pos, end)
        self._out_box.see(pos)
        self.root.after(2500, lambda: self._out_box.tag_remove(
            "hl_focus", "1.0", tk.END))
        self._status("已跳至查核問題位置")

    def _mark_issues_in_final(self):
        """Step 1：在最終稿件中用顏色標記所有查核問題。"""
        if not self._last_check_issues:
            messagebox.showinfo("無問題", "沒有待標記的查核問題")
            return
        final = self._out_box.get("1.0", tk.END).strip()
        if not final:
            messagebox.showwarning("最終稿件為空",
                                   "最終稿件尚無內容，請先生成報導")
            return
        # 先確保在主稿 tab
        if self._col3_active != "main":
            self._col3_switch_tab("main")
        # 清除舊標記
        self._out_box.tag_remove("hl_fixable", "1.0", tk.END)
        self._out_box.tag_remove("hl_unsourced", "1.0", tk.END)
        fixable_count = 0
        unsourced_count = 0
        for item in self._last_check_issues:
            val = item.get("value", "")
            if not val:
                continue
            typ = item.get("type", "")
            sug = item.get("suggestion", "")
            tag = "hl_fixable" if (typ in ("typo", "mismatch") and sug) else "hl_unsourced"
            start = "1.0"
            while True:
                pos = self._out_box.search(val, start, tk.END)
                if not pos:
                    break
                ep = f"{pos}+{len(val)}c"
                self._out_box.tag_add(tag, pos, ep)
                start = ep
                if tag == "hl_fixable":
                    fixable_count += 1
                else:
                    unsourced_count += 1
        parts = []
        if fixable_count:
            parts.append(f"{fixable_count} 處可修正（黃底）")
        if unsourced_count:
            parts.append(f"{unsourced_count} 處待人工確認（灰底）")
        self._status(f"已標記：{'、'.join(parts)}" if parts else "未找到匹配文字")

    def _auto_fix_issues(self):
        """Step 2：一鍵替換所有有 suggestion 的項目（typo + mismatch）。"""
        if not self._last_check_issues:
            return
        final = self._out_box.get("1.0", tk.END).strip()
        if not final:
            return
        # 先確保在主稿 tab
        if self._col3_active != "main":
            self._col3_switch_tab("main")
        fixed = 0
        for item in self._last_check_issues:
            val = item.get("value", "")
            sug = item.get("suggestion", "")
            typ = item.get("type", "")
            if not val or not sug or typ not in ("typo", "mismatch"):
                continue
            start = "1.0"
            while True:
                pos = self._out_box.search(val, start, tk.END)
                if not pos:
                    break
                ep = f"{pos}+{len(val)}c"
                self._out_box.delete(pos, ep)
                self._out_box.insert(pos, sug)
                start = f"{pos}+{len(sug)}c"
                fixed += 1
        # 清除已修正的標記
        self._out_box.tag_remove("hl_fixable", "1.0", tk.END)
        self._fix_chk_btn.configure(state="disabled", text="一鍵修正")
        self._status(f"已自動修正 {fixed} 處（灰底標記仍需人工確認）")

    # ── skill helpers ─────────────────────────────────────────────────────────
    def _update_skill_label(self):
        path = os.path.expanduser(self._skill_path_var.get())
        if os.path.exists(path):
            name = os.path.splitext(os.path.basename(path))[0]
            text, color = f"已載入  {name}", OK_FG
        else:
            text, color = "找不到技能包", ERR_FG
        lbl = self._skill_lbl
        if lbl is not None and lbl.winfo_exists():
            lbl.configure(text=text, text_color=color)

    def _browse_skill(self):
        path = filedialog.askopenfilename(
            title="選擇技能包",
            initialdir=os.path.join(SCRIPT_DIR, "skills"),
            filetypes=[("Markdown", "*.md"), ("所有", "*.*")])
        if not path: return
        self._skill_path_var.set(path)
        save_config({"skill_path": path})
        self._status(f"已載入技能包：{os.path.basename(path)}")

    # ══════════════════════════════════════════════════════════════════════════
    #  COLUMN 3 HANDLERS
    # ══════════════════════════════════════════════════════════════════════════
    def _update_charcount(self, _=None):
        if getattr(self, '_col3_active', 'main') == "side":
            chars = len(self._side_box.get("1.0", tk.END).strip())
            try: self._side_box.edit_modified(False)
            except Exception: pass
        else:
            chars = len(self._out_box.get("1.0", tk.END).strip())
            try: self._out_box.edit_modified(False)
            except Exception: pass
        self._charcount_lbl.configure(text=f"{chars} 字")

    def _clear_output(self):
        self._out_box.delete("1.0", tk.END)
        self._side_box.delete("1.0", tk.END)
        self._title_var.set("")
        self._side_title_var.set("")

    def _copy(self):
        parts = []
        t = self._title_var.get().strip()
        if t: parts.append(f"標題：{t}")
        parts.append(self._out_box.get("1.0", tk.END).strip())
        if self._sidebar_mode.get():
            st = self._side_title_var.get().strip()
            if st: parts.append(f"\n配稿標題：{st}")
            parts.append(self._side_box.get("1.0", tk.END).strip())
        text = "\n\n".join(p for p in parts if p)
        if not text: return
        self.root.clipboard_clear()
        self.root.clipboard_append(text)
        self._status("已複製至剪貼簿")

    def _save_txt(self):
        title = self._title_var.get().strip() or "記者會稿件"
        safe  = re.sub(r'[\\/*?:"<>|]', "", title)[:40]
        ts    = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        path  = filedialog.asksaveasfilename(
            defaultextension=".txt", initialfile=f"{safe}_{ts}.txt",
            filetypes=[("Text", "*.txt"), ("All", "*.*")])
        if not path: return
        author = self._author_var.get().strip()
        now    = datetime.datetime.now().isoformat()
        lines  = [
            f"標題：{self._title_var.get()}",
            f"作者：{author}",
            f"時間：{now}",
            "─" * 50, "",
            self._out_box.get("1.0", tk.END).strip()
        ]
        if self._sidebar_mode.get():
            lines += ["", "─" * 50,
                      f"配稿標題：{self._side_title_var.get()}", "",
                      self._side_box.get("1.0", tk.END).strip()]
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        self._status(f"已存：{os.path.basename(path)}")

    def _save_docx(self):
        if not HAS_DOCX:
            messagebox.showwarning("缺少套件",
                                    "請執行：pip install python-docx"); return
        from docx import Document as D
        from docx.shared import Pt
        title = self._title_var.get().strip() or "記者會稿件"
        safe  = re.sub(r'[\\/*?:"<>|]', "", title)[:40]
        ts    = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        path  = filedialog.asksaveasfilename(
            defaultextension=".docx", initialfile=f"{safe}_{ts}.docx",
            filetypes=[("Word", "*.docx"), ("All", "*.*")])
        if not path: return
        doc = D()
        doc.add_heading(title, level=1)
        doc.add_paragraph(
            f"{self._author_var.get()}  "
            f"{datetime.datetime.now().strftime('%Y/%m/%d %H:%M')}")
        doc.add_paragraph("─" * 30)
        for line in self._out_box.get("1.0", tk.END).splitlines():
            p = doc.add_paragraph(line)
            if p.runs: p.runs[0].font.size = Pt(12)
        if self._sidebar_mode.get():
            doc.add_heading(
                self._side_title_var.get() or "配稿", level=2)
            for line in self._side_box.get("1.0", tk.END).splitlines():
                p = doc.add_paragraph(line)
                if p.runs: p.runs[0].font.size = Pt(12)
        doc.save(path)
        self._status(f"已存 DOCX：{os.path.basename(path)}")

    def _status(self, msg: str):
        self.root.after(0, lambda: self._status_var.set(msg))


# ── entry point ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    ctk.set_appearance_mode("dark")
    ctk.set_default_color_theme("dark-blue")
    root = ctk.CTk()
    PressConfStudio(root)
    root.mainloop()
